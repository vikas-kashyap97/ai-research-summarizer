from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
import streamlit as st
import fitz as ft
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from gtts import gTTS
import base64



# Load API keys from .env file
load_dotenv()

# Page setup
st.set_page_config(page_title="AI Research Summarizer", layout="wide")
st.header('📄 AI-Powered Research Summarizer')

# Initialize model
model = ChatGoogleGenerativeAI(model='gemini-1.5-pro-latest')

# UI Dropdowns
paper_input = st.selectbox("📘 Select type of paper:", 
                           ["Research Article", "Review Paper", "Technical Report", "Thesis", "Whitepaper", "Conference Paper"])

style_input = st.selectbox("📝 Select summary style:", 
                           ["Beginner Friendly", "Technical", "Code Oriented", "Mathematical", "Academic"])

length_input = st.selectbox("📏 Select summary length:", [
    "Short (1-2 paragraphs)",
    "Medium (3-5 paragraphs)",
    "Long (Detailed explanation with subpoints)",
    "Extended (Section-wise breakdown with examples and analogies)",
    "Comprehensive (In-depth explanation with mathematics, context, and practical implications)"
])
language_input = st.selectbox("🌍 Select output language:", [
    "English", "Hindi", "Spanish", "French", "German", "Chinese", "Arabic", "Portuguese", "Russian", "Japanese"
])


uploaded_file = st.file_uploader("📤 Upload a research paper (PDF)", type=["pdf"])

extracted_text = ""

if uploaded_file is not None:
    with ft.open(stream=uploaded_file.read(), filetype="pdf") as doc:
        for page in doc:
            extracted_text += page.get_text()
    
    st.success("✅ PDF parsed successfully!")

user_input = st.text_area("📚 Paste the abstract or full content of the research paper here (or use uploaded PDF):", value=extracted_text)


# Define the Prompt Template
prompt_template = PromptTemplate(
    input_variables=["paper_type", "summary_style", "summary_length", "paper_content", "output_language"],
    template="""
You are a world-class academic summarization assistant with expertise in scientific, technical, and mathematical domains. 
Your role is to create detailed, accurate, and well-structured summaries of complex research content tailored to various audiences and comprehension levels.

**TASK OVERVIEW:**
Summarize the following *{paper_type}* in a **{summary_style}** style and at a **{summary_length}** detail level. 

**YOUR SUMMARY SHOULD:**
- Be highly accurate and faithful to the original research paper content.
- Clearly outline the objective of the research, methodology, key findings, and their implications.
- Use appropriate terminology, formatting, and domain-specific expressions based on the selected summary style.
- If the selected style is *Mathematical*, provide equations, proofs, or quantitative results where applicable.
- If the style is *Code Oriented*, include relevant pseudocode, programming logic, or computational flow.
- If *Beginner Friendly*, break down technical terms, use analogies, and simplify concepts without losing core meaning.
- For *Academic* or *Technical* styles, maintain a formal tone and include references to methodologies, prior research, and detailed analysis.
- Support the explanation with real-world analogies or comparisons to make complex topics more intuitive.
- Include visual or structural segmentation cues if applicable (e.g., bullet points, subheadings, key takeaway boxes).

**FOR EXTENDED OR COMPREHENSIVE SUMMARIES:**
- Include section-wise breakdowns (e.g., Abstract, Introduction, Methodology, Results, Discussion, Conclusion).
- Discuss the significance of the study in the context of its domain or industry.
- Highlight theoretical foundations, mathematical models, or simulations used.
- Include any limitations, future scope, and potential real-world applications.
- Maintain academic integrity, neutrality, and objective representation of the content.

---

**Research Paper Content:**

\"\"\"{paper_content}\"\"\"

---
Respond in the **{output_language}** language.

Now generate the summary based on the above instructions and selected configurations:

- Paper Type: {paper_type}
- Style: {summary_style}
- Summary Length: {summary_length}
"""
)

# Generate and run the prompt
if st.button('🔍 Generate Summary'):
    if user_input.strip() == "":
        st.warning("⚠️ Please provide the research paper content before summarizing.")
    else:
        prompt = prompt_template.format(
            paper_type=paper_input,
            summary_style=style_input,
            summary_length=length_input,
            paper_content=user_input,
            output_language=language_input
        )

        result = model.invoke(prompt)
        summary_text = result.content

        st.subheader("📑 Summary Output")
        st.write(summary_text)

                # ✅ TTS (gTTS browser-based)
        def generate_tts_audio(text):
            tts = gTTS(text)
            buffer = BytesIO()
            tts.write_to_fp(buffer)
            buffer.seek(0)
            return buffer

        if st.button("🔊 Read Summary Out Loud"):
            audio_buffer = generate_tts_audio(summary_text)
            audio_base64 = base64.b64encode(audio_buffer.read()).decode()
            audio_html = f"""
                <audio autoplay controls>
                    <source src="data:audio/mp3;base64,{audio_base64}" type="audio/mp3">
                    Your browser does not support the audio element.
                </audio>
            """
            st.markdown(audio_html, unsafe_allow_html=True)

        # ✅ DOCX EXPORT
        def generate_docx(text):
            buffer = BytesIO()
            doc = Document()
            doc.add_heading("AI Research Summary", level=0)
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            for para in text.split("\n"):
                doc.add_paragraph(para.strip())
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        # ✅ PDF EXPORT
        def generate_pdf(text):
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = [Paragraph("AI Research Summary", styles['Title']), Spacer(1, 12)]
            for para in text.split("\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['BodyText']))
                    story.append(Spacer(1, 8))
            doc.build(story)
            buffer.seek(0)
            return buffer

        # ✅ EXPORT BUTTONS
        st.markdown("### 📤 Export Summary")
        col1, col2 = st.columns(2)

        with col1:
            docx_file = generate_docx(summary_text)
            st.download_button(
                label="💾 Download as DOCX",
                data=docx_file,
                file_name="summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col2:
            pdf_file = generate_pdf(summary_text)
            st.download_button(
                label="📄 Download as PDF",
                data=pdf_file,
                file_name="summary.pdf",
                mime="application/pdf"
            )
